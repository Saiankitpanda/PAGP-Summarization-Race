from docx import Document
import os 


def generate_sample():
    doc = Document()
    doc.add_heading("SERVICE AGREEMENT", 0)
    doc.add_paragraph("This AGREEMENT is made on 2024-01-29.")
    doc.add_paragraph(
        'BETWEEN: ACME CORP (hereinafter "Company") AND JOHN DOE (hereinafter "Contractor").'
    )

    # Add enough text to ensure chunking happens (aim for > 6000 chars)
    # 6000 chars is roughly 1000-1200 words.

    dummy_text = (
        "The Contractor agrees to perform the following services in a professional manner. "
        * 50
    )
    doc.add_paragraph(dummy_text)

    for i in range(10):
        doc.add_heading(f"Article {i + 1}", level=1)
        doc.add_paragraph(f"Details regarding Article {i + 1}. " + dummy_text)

    output_path = "/home/jarvis/Desktop/PAGP/legal-templater/input/sample.docx"
    doc.save(output_path)
    print(f"Generated sample doc at {output_path}")


if __name__ == "__main__":
    generate_sample()
